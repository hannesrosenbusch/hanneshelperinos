# pylint: disable=C0301, C0103, C0121, C0209, W0212, R*
'''survey exporter input and output function (originally authored by Tim Wienboeker)'''

import time
import json
from datetime import date
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH # pylint: disable=E0611
from docx.shared import RGBColor
import requests
from ipywidgets import interact
import pkg_resources


def user_input():
    '''get user inputs from widget interface in colan notebook
    inputs consist of long jsonstring with survey architecture, and two booleans'''
    inputs = [None, None, None]


    def returner(Survey):
        inputs[0] = Survey

    def returner1(include_images):
        inputs[1] = include_images

    def returner2(english_labels):
        inputs[2] = english_labels
    interact(returner, Survey="Paste big thing here...")
    interact(returner1, include_images=False)
    interact(returner2, english_labels=False)
    return inputs


def go(inputs):
    '''process inputs to generate word file displaying the survey archtecture'''
    t = inputs[0]
    include_images = inputs[1]
    english_lang = inputs[2]

    # dictionary of question types:
    if english_lang:
        dict_qtypes = {"mc": "Multiple Choice",
                       "freetext": "Open question",
                       "info": "Info box",
                       "matrix": "Matrix",
                       "likert": "Likert",
                       "imagecloud": "Multiple Choice (with images)",
                       "image": "Multiple Choice (with images and text)",
                       "numericslider": "Numeric slider / NPS",
                       "ranking": "Ranking",
                       "starslider": "Stars",
                       "propertyslider": "Preference slider",
                       "number": "Number (Open entry)",
                       "heatmap": "Heatmap",
                       "videoplay": "Audio/Video",
                       "photocaptur": "Take photo"}
    else:
        dict_qtypes = {"mc": "Multiple Choice",
                       "freetext": "Offene Frage",
                       "info": "Infobox",
                       "matrix": "Matrix",
                       "likert": "Likert",
                       "numericslider": "Numerischer Slider / NPS",
                       "ranking": "Ranking",
                       "starslider": "Stars",
                       "propertyslider": "Präferenz-Slider",
                       "number": "Zahl (Freie Eingabe)",
                       "heatmap": "Heatmap",
                       "videoplay": "Audio/Video",
                       "photocaptur": "Fotoaufnahme"}

    json_string = t.rstrip()
    survey_raw = json.loads(json_string)
    document = Document()

    # General style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Roboto'
    font.size = Pt(10)
    paragraph_format = style.paragraph_format

    # Add header for first page with table (logo on the left, address on the right)
    header = document.sections[0].first_page_header
    document.sections[0].different_first_page_header_footer = True
    htable = header.add_table(1, 4, Inches(6.25))
    htab_cells = htable.rows[0].cells
    ht0 = htab_cells[0].paragraphs[0]  # cell including the logo
    kh = ht0.add_run(style=None)
    stream = pkg_resources.resource_stream(__name__, 'data/Appinio-Logo.png')
    kh.add_picture(stream, width=Inches(1.401575))

    # cell including address and contact information
    ht1 = htab_cells[3].paragraphs[0]
    run = ht1.add_run(
        "APPINIO GmbH\nGroße Theaterstraße31\n20354 Hamburg\n\ncontact@appinio.com\n+49 40 / 413 49 710\nwww.appinio.com")
    run.font.name = "Roboto"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(69, 107, 132)
    ht1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # header for following pages only with Appinio logo
    header = document.sections[0].header
    document.sections[0].different_first_page_header_footer = True
    htable = header.add_table(2, 2, Inches(6))
    htab_cells = htable.rows[0].cells
    ht0 = htab_cells[0].paragraphs[0]  # cell including the logo
    kh = ht0.add_run(style=None)
    kh.add_picture(stream, width=Inches(1.401575))

    # Add title and date
    document.add_paragraph()
    para = document.add_paragraph()
    run = para.add_run('\t\t Hamburg, ' +
                       str(date.today().strftime("%d.%m.%Y")))
    run.font.color.rgb = RGBColor(69, 107, 132)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph()

    run = document.add_paragraph().add_run(survey_raw['title'])
    font = run.font
    font.name = 'Roboto Medium'
    font.size = Pt(14)
    font.color.rgb = RGBColor(5, 49, 73)
    document.add_paragraph()

    # add table
    table = document.add_table(1, 3)
    table.style = 'TableGrid'

    # populate header row
    heading_cells = table.rows[0].cells
    if english_lang:
        heading_cells[0].paragraphs[0].add_run('Question no.').bold = True
        heading_cells[1].paragraphs[0].add_run('Survey').bold = True
        heading_cells[2].paragraphs[0].add_run('Question type').bold = True
    else:
        heading_cells[0].paragraphs[0].add_run('Frage').bold = True
        heading_cells[1].paragraphs[0].add_run('Fragebogen').bold = True
        heading_cells[2].paragraphs[0].add_run('Fragetyp').bold = True

    # contents
    info_no = 1
    question_no = 1
    hidden = []

    for i in range(len(survey_raw['questions'])):  # iterate over questions
        if survey_raw['questions'][i]['hideForCompany']:
            hidden.append(i)
            continue
        cells = table.add_row().cells

        # cell 1 -------------
        if survey_raw['questions'][i]['qtype'] == 'info' or survey_raw['questions'][i]['qtype'] == 'videoplay':
            cells[0].paragraphs[0].add_run("Info "+str(info_no)).bold = True
            info_no += 1
        else:
            cells[0].paragraphs[0].add_run("F"+str(question_no)).bold = True
            question_no += 1

        # cell 2 -------------
        # question wording -------------
        cells[1].paragraphs[0].add_run(survey_raw['questions'][i]['text']).bold = True
        paragraph_format.space_before = Pt(2)
        paragraph_format.space_after = Pt(2)

        # Include instructions for participants
        # check whether instructions for participants exist
        if "help" in survey_raw['questions'][i]:
            cells[1].add_paragraph(str(survey_raw['questions'][i]['help']))

        # Include image
        if include_images:
            # check whether image URL exist
            if "media" in survey_raw['questions'][i]:
                image_url = str(survey_raw['questions'][i]['media'])
                response = requests.get(image_url)
                binary_img = BytesIO(response.content)
                paragraph = cells[1].paragraphs[0]
                run = paragraph.add_run()
                run.add_break()
                run.add_picture(binary_img, width=Inches(2))
                time.sleep(1)

        cells[1].add_paragraph()

        # Answers -------------
        if survey_raw['questions'][i]['qtype'] == 'matrix':
            if english_lang:
                cells[1].add_paragraph("Answers:").bold = True
            else:
                cells[1].add_paragraph("Antworten:").bold = True

        # Answer text and letter
        answer_letter = 65  # because chr(65) = 'A'

        # Check for randomization
        random_answer = False
        random_answer_text = " "
        for j in range(len(survey_raw['questions'][i]['answers'])):
            if "random" in survey_raw['questions'][i]['answers'][j]:
                if survey_raw['questions'][i]['answers'][j]["random"] == True:
                    random_answer = True

        # If randomized:
        if random_answer:
            for j in range(len(survey_raw['questions'][i]['answers'])):  # iterate over answers if randomized
                if not survey_raw['questions'][i]['answers'][j]["random"]:
                    if english_lang:
                        random_answer_text = " (not randomized)"
                    else:
                        random_answer_text = " (nicht randomisiert)"

                answer_filter_text = " "
                # if include_filters:

                #     # Filter IF
                #     if survey_raw['questions'][i]['answers'][j]["filterRequirements"] != []:
                #         for k in range(len(survey_raw['questions'][i]['answers'][j]["filterRequirements"])):
                #             filter_id = str(
                #                 survey_raw['questions'][i]['answers'][j]["filterRequirements"][k])
                #             for l in range(len(survey_raw['questions'])):
                #                 for m in range(len(survey_raw['questions'][l]['answers'])):
                #                     if "filterId" in survey_raw['questions'][l]['answers'][m].keys():
                #                         if filter_id in survey_raw['questions'][l]['answers'][m]["filterId"]:
                #                             infoxboxes_until_here = info_no - 1
                #                             filter_question_no = 1 + l - infoxboxes_until_here
                #                             filter_answer_letter = 65 + m
                #                             answer_filter_text = str(
                #                                 " (IF F"+str(filter_question_no)+str(chr(filter_answer_letter))+")"+random_answer_text)
                #                 for n in range(len(survey_raw['questions'][l]['key'])):
                #                     if filter_id in survey_raw['questions'][l]['key'][n]["filterId"]:
                #                         infoxboxes_until_here = info_no - 1
                #                         filter_question_no = 1 + l - infoxboxes_until_here
                #                         filter_answer_letter = 65 + n
                #                         answer_filter_text = str(
                #                             " (IF F"+str(filter_question_no)+str(chr(filter_answer_letter))+")"+random_answer_text)

                #     # Filter IF NOT
                #     if survey_raw['questions'][i]['answers'][j]["filterNotRequirements"] != []:
                #         for k in range(len(survey_raw['questions'][i]['answers'][j]["filterNotRequirements"])):
                #             filter_id = str(
                #                 survey_raw['questions'][i]['answers'][j]["filterNotRequirements"][k])
                #             for l in range(len(survey_raw['questions'])):
                #                 for m in range(len(survey_raw['questions'][l]['answers'])):
                #                     if "filterId" in survey_raw['questions'][l]['answers'][m].keys():
                #                         if filter_id in survey_raw['questions'][l]['answers'][m]["filterId"]:
                #                             infoxboxes_until_here = info_no - 1
                #                             filter_question_no = 1 + l - infoxboxes_until_here
                #                             filter_answer_letter = 65 + m
                #                             answer_filter_text = str(
                #                                 " (IF NOT F"+str(filter_question_no)+str(chr(filter_answer_letter))+")"+random_answer_text)
                #                 for n in range(len(survey_raw['questions'][l]['key'])):
                #                     if filter_id in survey_raw['questions'][l]['key'][n]["filterId"]:
                #                         infoxboxes_until_here = info_no - 1
                #                         filter_question_no = 1 + l - infoxboxes_until_here
                #                         filter_answer_letter = 65 + n
                #                         answer_filter_text = str(
                #                             " (IF NOT F"+str(filter_question_no)+str(chr(filter_answer_letter))+")"+random_answer_text)

                answer_text = str(chr(
                    answer_letter) + ": " + survey_raw['questions'][i]['answers'][j]['text'] + answer_filter_text+random_answer_text)
                cells[1].add_paragraph(answer_text)
                answer_letter += 1

        if not random_answer:
            for j in range(len(survey_raw['questions'][i]['answers'])):  # iterate over answers if not randomized
                answer_filter_text = " "
                # if include_filters:

                #     # Filter IF
                #     if survey_raw['questions'][i]['answers'][j]["filterRequirements"] != []:
                #         for k in range(len(survey_raw['questions'][i]['answers'][j]["filterRequirements"])):
                #             filter_id = str(
                #                 survey_raw['questions'][i]['answers'][j]["filterRequirements"][k])
                #             for l in range(len(survey_raw['questions'])):
                #                 for m in range(len(survey_raw['questions'][l]['answers'])):
                #                     if "filterId" in survey_raw['questions'][l]['answers'][m].keys():
                #                         if filter_id in survey_raw['questions'][l]['answers'][m]["filterId"]:
                #                             if survey_raw['questions'][i]['qtype'] == 'info' or survey_raw['questions'][i]['qtype'] == 'videoplay':
                #                                 infoxboxes_until_here = info_no
                #                             else:
                #                                 infoxboxes_until_here = info_no - 1
                #                             filter_question_no = 1 + l - infoxboxes_until_here
                #                             filter_answer_letter = 65 + m
                #                             answer_filter_text = str(
                #                                 " (IF F"+str(filter_question_no)+str(chr(filter_answer_letter))+")")
                #                 for n in range(len(survey_raw['questions'][l]['key'])):
                #                     if filter_id in survey_raw['questions'][l]['key'][n]["filterId"]:
                #                         if survey_raw['questions'][i]['qtype'] == 'info' or survey_raw['questions'][i]['qtype'] == 'videoplay':
                #                             infoxboxes_until_here = info_no
                #                         else:
                #                             infoxboxes_until_here = info_no - 1
                #                         filter_question_no = 1 + l - infoxboxes_until_here
                #                         filter_answer_letter = 65 + n
                #                         answer_filter_text = str(
                #                             " (IF F"+str(filter_question_no)+str(chr(filter_answer_letter))+")")

                #     # Filter IF NOT
                #     if survey_raw['questions'][i]['answers'][j]["filterNotRequirements"] != []:
                #         for k in range(len(survey_raw['questions'][i]['answers'][j]["filterNotRequirements"])):
                #             filter_id = str(
                #                 survey_raw['questions'][i]['answers'][j]["filterNotRequirements"][k])
                #             for l in range(len(survey_raw['questions'])):
                #                 for m in range(len(survey_raw['questions'][l]['answers'])):
                #                     if "filterId" in survey_raw['questions'][l]['answers'][m].keys():
                #                         if filter_id in survey_raw['questions'][l]['answers'][m]["filterId"]:
                #                             infoxboxes_until_here = info_no - 1
                #                             filter_question_no = 1 + l - infoxboxes_until_here
                #                             filter_answer_letter = 65 + m
                #                             answer_filter_text = str(
                #                                 " (IF NOT F"+str(filter_question_no)+str(chr(filter_answer_letter))+")")
                #                 for n in range(len(survey_raw['questions'][l]['key'])):
                #                     if filter_id in survey_raw['questions'][l]['key'][n]["filterId"]:
                #                         infoxboxes_until_here = info_no - 1
                #                         filter_question_no = 1 + l - infoxboxes_until_here
                #                         filter_answer_letter = 65 + n
                #                         answer_filter_text = str(
                #                             " (IF NOT F"+str(filter_question_no)+str(chr(filter_answer_letter))+")")

                answer_text = str(chr(answer_letter)+": "+survey_raw['questions'][i]['answers'][j]['text'] + answer_filter_text)
                cells[1].add_paragraph(answer_text)
                answer_letter += 1


        #add optional custom text entry
        if "allowCustomText" in survey_raw['questions'][i]:
            try:
                if survey_raw['questions'][i]["allowCustomText"]:
                    freetex_answer_test = survey_raw['questions'][i]["customTextName"]
                    if not survey_raw['questions'][i]['answers'][j]["random"]:
                        if english_lang:
                            cells[1].add_paragraph(
                                str(str(freetex_answer_test) + " (Freetext)"))
                        else:
                            cells[1].add_paragraph(
                                str(str(freetex_answer_test) + " (Freitext)"))
            except: # pylint: disable=W0702
                pass

        # info texts
        if survey_raw['questions'][i]['qtype'] == 'info' or survey_raw['questions'][i]['qtype'] == 'videoplay':
            try:
                cells[1].add_paragraph(
                    str(survey_raw['questions'][i]['infoText']))
            except: # pylint: disable=W0702
                pass

        # include images
        for j in range(len(survey_raw['questions'][i]['answers'])):
            if include_images:
                # check whether image URL exist
                if "imageUrl" in survey_raw['questions'][i]['answers'][j]:
                    image_url = str(
                        survey_raw['questions'][i]['answers'][j]['imageUrl'])
                    print(image_url)
                    response = requests.get(image_url)
                    binary_img = BytesIO(response.content)
                    paragraph = cells[1].paragraphs[0]
                    run = paragraph.add_run()
                    run.add_break()
                    run.add_picture(binary_img, width=Inches(2))
                    time.sleep(1)

        # matrix scale items
        random_mat_items = False
        random_mat_items_text = " "
        for j in range(len(survey_raw['questions'][i]['rows'])):
            if "random" in survey_raw['questions'][i]['rows'][j]:
                if survey_raw['questions'][i]['rows'][j]["random"]:
                    random_mat_items = True

        # if randomized
        if random_mat_items:
            for j in range(len(survey_raw['questions'][i]['rows'])):  # answers
                if not survey_raw['questions'][i]['rows'][j]["random"]:
                    if english_lang:
                        random_mat_items_text = " (not randomized)"
                    else:
                        random_mat_items_text = " (nicht randomisiert)"

            answer_letter = 65  # because chr(65) = 'A'
            if survey_raw['questions'][i]['rows'] != []:  # Test whether matrix items exist
                cells[1].add_paragraph()
                cells[1].add_paragraph("Items:")
                # matrix items
                for k in range(len(survey_raw['questions'][i]['rows'])):
                    item_text = str(chr(
                        answer_letter) + ": "+survey_raw['questions'][i]['rows'][k]['text'] + random_mat_items_text)
                    cells[1].add_paragraph(item_text)
                    answer_letter += 1

        # if not randomized
        if not random_mat_items:
            answer_letter = 65  # because chr(65) = 'A'
            if survey_raw['questions'][i]['rows'] != []:  # Test whether matrix items exist
                cells[1].add_paragraph()
                cells[1].add_paragraph("Items:")
                # matrix items
                for k in range(len(survey_raw['questions'][i]['rows'])):
                    item_text = str(chr(answer_letter)+": " +
                                    survey_raw['questions'][i]['rows'][k]['text'])
                    cells[1].add_paragraph(item_text)
                    answer_letter += 1

        # likert scale
        answer_letter = 65  # because chr(65) = 'A'
        if survey_raw['questions'][i]['key'] != []:  # Test whether matrix items exist
            for k in range(len(survey_raw['questions'][i]['key'])):  # matrix items
                item_text = str(chr(answer_letter) + ": " +
                                survey_raw['questions'][i]['key'][k]['text'])
                cells[1].add_paragraph(item_text)
                answer_letter += 1

        # cell 3 -------------
        # "Fragetyp" -------------

        if survey_raw['questions'][i]['qtype'] != "mc":
            for abbreviation, new_label in dict_qtypes.items():
                if abbreviation == survey_raw['questions'][i]['qtype']:
                    cells[2].paragraphs[0].add_run(new_label)
        else:
            if 'multioptions' in survey_raw['questions'][i]:
                if survey_raw['questions'][i]['multioptions']:
                    cells[2].paragraphs[0].add_run("Multiple Choice")
                else:
                    cells[2].paragraphs[0].add_run("Single Choice")

        # Randomisation
        if random_mat_items:
            if english_lang:
                cells[2].add_paragraph("(Items randomized)")
            else:
                cells[2].add_paragraph("(Items randomisiert)")

        if random_answer:
            if english_lang:
                cells[2].add_paragraph("(Answers randomized)")
            else:
                cells[2].add_paragraph("(Antworten randomisiert)")

        # Max options
        if "maxOptions" in survey_raw['questions'][i]:
            if english_lang:
                cells[2].add_paragraph(
                    "Max Answers: "+str(survey_raw['questions'][i]['maxOptions']))
            else:
                cells[2].add_paragraph(
                    "Max Antworten: "+str(survey_raw['questions'][i]['maxOptions']))

        # Min options
        if "minOptions" in survey_raw['questions'][i]:
            if english_lang:
                cells[2].add_paragraph(
                    "Min Answers: "+str(survey_raw['questions'][i]['minOptions']))
            else:
                cells[2].add_paragraph(
                    "Min Antworten: "+str(survey_raw['questions'][i]['minOptions']))

        # Do filter exist?
        if survey_raw['questions'][i]["filterRequirements"] != [] or survey_raw['questions'][i]["filterNotRequirements"] != []:
            cells[2].add_paragraph()
            cells[2].add_paragraph("Filter:")

        # Filter IF
        if survey_raw['questions'][i]["filterRequirements"] != []: #if the current question has filter requirements
            for k in range(len(survey_raw['questions'][i]["filterRequirements"])): #iterate over filter requirements
                filter_id = str(survey_raw['questions'][i]["filterRequirements"][k]) #select current filter
                for l in range(len(survey_raw['questions'])): #iterate over questions2
                    if survey_raw['questions'][l]['answers'] != []: #if current question2 has answers
                        for m in range(len(survey_raw['questions'][l]['answers'])): #iterate over answers
                            if "filterId" in survey_raw['questions'][l]['answers'][m].keys(): #if current answer has filterId
                                if filter_id in survey_raw['questions'][l]['answers'][m]["filterId"]: #if filter of current answer (question2) pertains to original question
                                    if survey_raw['questions'][l]['qtype'] == 'info' or survey_raw['questions'][l]['qtype'] == 'videoplay':
                                        infoxboxes_until_here = info_no
                                    else:
                                        infoxboxes_until_here = info_no - 1
                                    print("Infoboxes:" + str(infoxboxes_until_here))
                                    print("l:"+str())
                                    filter_question_no = 1 + l - len(hidden)
                                    print("filter_qestion:" +
                                          str(filter_question_no))
                                    filter_answer_letter = 65 + m
                                    cells[2].add_paragraph("IF F" + str(filter_question_no) + str(chr(filter_answer_letter)))
                    elif survey_raw['questions'][l]['key'] != []:
                        for n in range(len(survey_raw['questions'][l]['key'])):
                            if filter_id in survey_raw['questions'][l]['key'][n]["filterId"]:
                                # if survey_raw['questions'][l]['qtype'] == 'info' or survey_raw['questions'][l]['qtype'] == 'videoplay':
                                #     infoxboxes_until_here = info_no
                                # else:
                                #     infoxboxes_until_here = info_no - 1
                                filter_question_no = 1 + l  - len(hidden) #- infoxboxes_until_here
                                filter_answer_letter = 65 + n
                                cells[2].add_paragraph("IF F" + str(filter_question_no) + str(chr(filter_answer_letter)))

        # Filter IF NOT
        elif survey_raw['questions'][i]["filterNotRequirements"] != []:
            for k in range(len(survey_raw['questions'][i]["filterNotRequirements"])):
                filter_id = str(survey_raw['questions']
                                [i]["filterNotRequirements"][k])
                for l in range(len(survey_raw['questions'])):
                    if survey_raw['questions'][l]['answers'] != []:
                        for m in range(len(survey_raw['questions'][l]['answers'])):
                            if "filterId" in survey_raw['questions'][l]['answers'][m].keys():
                                if filter_id in survey_raw['questions'][l]['answers'][m]["filterId"]:
                                    #infoxboxes_until_here = info_no - 1
                                    filter_question_no = 1 + l  - len(hidden)#- infoxboxes_until_here
                                    filter_answer_letter = 65 + m
                                    cells[2].add_paragraph("IF NOT F"+str(filter_question_no) + str(chr(filter_answer_letter)))
                    if survey_raw['questions'][l]['key'] != []:
                        for n in range(len(survey_raw['questions'][l]['key'])):
                            if filter_id in survey_raw['questions'][l]['key'][n]["filterId"]:
                                #infoxboxes_until_here = info_no - 1
                                filter_question_no = 1 + l  - len(hidden)#- infoxboxes_until_here
                                filter_answer_letter = 65 + n
                                cells[2].add_paragraph("IF NOT F" + str(filter_question_no) + str(chr(filter_answer_letter)))

    # Layout stuff
    for row in table.rows:
        for cell, width in zip(row.cells, (Inches(0.72), Inches(5.1), Inches(1.50))):
            cell.width = width

    # Set a cell background (shading) color to RGB D9D9D9.
    shading1 = parse_xml(r'<w:shd {} w:fill="053149"/>'.format(nsdecls('w')))
    table.cell(0, 0)._tc.get_or_add_tcPr().append(shading1)
    shading2 = parse_xml(r'<w:shd {} w:fill="053149"/>'.format(nsdecls('w')))
    table.cell(0, 1)._tc.get_or_add_tcPr().append(shading2)
    shading3 = parse_xml(r'<w:shd {} w:fill="053149"/>'.format(nsdecls('w')))
    table.cell(0, 2)._tc.get_or_add_tcPr().append(shading3)

    # Save it
    file_name = survey_raw['title'].replace(":", "-")
    file_name = file_name.replace("/", "-")
    document.save(str(file_name + '_export.docx'))
